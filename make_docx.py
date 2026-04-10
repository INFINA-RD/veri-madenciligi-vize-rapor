"""
PDF'teki icerigi ayni sekilde Word dosyasina donusturen yardimci betik.
Mevcut notebook, veri ve modeller ile ETKILESIME GIRMEZ: sadece outputs/
klasorundeki hazir gorselleri ve Proje_Raporu.pdf'ten elde edilen metni
kullanarak 'Proje_Raporu.docx' dosyasini olusturur.
"""

import os
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE, "outputs")
DOCX_PATH = os.path.join(BASE, "Proje_Raporu.docx")


def set_columns(section, num=2, space_cm=0.6):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        col = cols[0]
    else:
        col = OxmlElement("w:cols")
        sectPr.append(col)
    col.set(qn("w:num"), str(num))
    col.set(qn("w:space"), str(int(space_cm * 567)))  # 1 cm ~ 567 twip
    col.set(qn("w:equalWidth"), "1")


def add_heading(doc, text, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                italic=False, space_before=8, space_after=4, small_caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if small_caps:
        rPr = r._element.get_or_add_rPr()
        sc = OxmlElement("w:smallCaps")
        sc.set(qn("w:val"), "1")
        rPr.append(sc)
    return p


def add_body(doc, text, size=10, justify=True, first_line_cm=0.5,
             space_after=2, italic=False, bold=False,
             align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if first_line_cm:
        pf.first_line_indent = Cm(first_line_cm)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r.italic = italic
    r.bold = bold
    return p


def add_figure(doc, path, caption, width_cm=8.2):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.name = "Times New Roman"


def build():
    doc = Document()

    # Sayfa kenar bosluklari (IEEE-like)
    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    # Varsayilan stil
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # ---- BASLIK (tek sutun) ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run(
        "ABD Borsasi Finansal Verileri Uzerinde Makine Ogrenmesi\n"
        "Tabanli Sirket Gelir Tahmini: Bir Regresyon Yaklasimi"
    )
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Times New Roman"

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(2)
    ar = author.add_run("Ogrenci Adi Soyadi")
    ar.font.size = Pt(11)
    ar.font.name = "Times New Roman"

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(2)
    ir = info.add_run("Veri Madenciligi Dersi Vize Projesi")
    ir.font.size = Pt(10)
    ir.italic = True
    ir.font.name = "Times New Roman"

    info2 = doc.add_paragraph()
    info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info2.paragraph_format.space_after = Pt(10)
    ir2 = info2.add_run("Bilgisayar Muhendisligi Bolumu")
    ir2.font.size = Pt(10)
    ir2.italic = True
    ir2.font.name = "Times New Roman"

    # ---- 2 SUTUNLU BOLUMU BASLAT ----
    new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    new_sec.top_margin = Cm(1.9)
    new_sec.bottom_margin = Cm(1.9)
    new_sec.left_margin = Cm(1.6)
    new_sec.right_margin = Cm(1.6)
    set_columns(new_sec, num=2, space_cm=0.6)

    # ---- OZET ----
    abs_p = doc.add_paragraph()
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.space_after = Pt(4)
    run = abs_p.add_run("Ozet—")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    run2 = abs_p.add_run(
        "Bu calismada, 2014-2018 yillarini kapsayan ve 4000 civari ABD borsasi sirketinin "
        "224 finansal gostergesini iceren bir veri seti kullanilarak, bir sirketin bir sonraki "
        "yilki yillik gelirinin (Revenue) tahmin edilmesine yonelik bir regresyon modeli "
        "gelistirilmistir. Ardisik yil ciftleri (2014-2015, 2015-2016, 2016-2017, 2017-2018) "
        "hisse sembolu uzerinden birlestirilerek yaklasik 15.500 egitim ornegi elde edilmistir. "
        "Hedef degiskenin uzun kuyruklu dagilimindan oturu logaritmik donusum (log1p) "
        "uygulanmistir. Veri temizleme, eksik deger atama, one-hot kodlama, standartlastirma "
        "ve ANOVA F-istatistigi tabanli ozellik secimi adimlarinin ardindan yedi farkli "
        "regresyon algoritmasi (Dogrusal, Ridge, Lasso, Karar Agaci, Rastgele Orman, Gradient "
        "Boosting ve K-En Yakin Komsu) 5-katli capraz dogrulama ile karsilastirilmistir. "
        "Deneyler, Gradient Boosting modelinin log olcekli test kumesinde 0.966 R2 ve 0.228 MAE "
        "ile en iyi performansi sagladigini gostermistir. Bulgular, finansal tablo "
        "gostergelerinin kisa vadeli gelir tahmini icin guclu bir sinyal tasidigini ve topluluk "
        "yontemlerinin dogrusal yontemlere gore belirgin ustunluk sagladigini ortaya koymaktadir."
    )
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = "Times New Roman"

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.paragraph_format.space_after = Pt(8)
    kwr1 = kw.add_run("Anahtar Kelimeler—")
    kwr1.bold = True
    kwr1.italic = True
    kwr1.font.size = Pt(9)
    kwr1.font.name = "Times New Roman"
    kwr2 = kw.add_run(
        "veri madenciligi, regresyon, gelir tahmini, makine ogrenmesi, gradient boosting, "
        "rastgele orman, ozellik muhendisligi, k-katli capraz dogrulama."
    )
    kwr2.italic = True
    kwr2.font.size = Pt(9)
    kwr2.font.name = "Times New Roman"

    # ---- I. GIRIS ----
    add_heading(doc, "I. GIRIS", size=10, small_caps=True)
    add_body(doc,
        "Bir sirketin gelecekteki gelirini (Revenue) dogru bicimde tahmin edebilmek; "
        "stratejik planlama, kaynak tahsisi, kredi degerlendirmesi ve yatirimci raporlamasi "
        "gibi bircok alanda kritik oneme sahiptir. Geleneksel finansal tahmin yontemleri "
        "genellikle uzman bilgisi, oran analizi ve sinirli sayida degiskene dayanmaktadir. "
        "Oysa gunumuzde halka acik sirketlerin yuzlerce finansal gostergesi duzenli olarak "
        "raporlanmakta; bu verilerin butun halinde kullanimi ise makine ogrenmesi yontemleri "
        "olmadan zordur.")
    add_body(doc,
        "Bu calismanin amaci, 2014-2018 yillarini kapsayan 200+ finansal gostergelik bir "
        "veri seti kullanarak, her sirket icin bir sonraki yilki geliri regresyon yontemleriyle "
        "tahmin eden bir veri madenciligi sureci gelistirmektir. Problem denetimli ogrenme "
        "cercevesinde regresyon olarak tanimlanmistir: surekli degerli hedef, bir sonraki "
        "yilin gelir miktaridir. Calismada ardisik yil ciftleri hisse sembolu uzerinden "
        "eslestirilmis; yil N ozellikleri ile yil N+1 geliri arasinda bir eslestirme yapilmistir.")
    add_body(doc,
        "Raporun geri kalani su sekilde duzenlenmistir: Bolum II veri seti ve kullanilan "
        "yontemleri aciklar. Bolum III deneysel sonuclari sunar. Bolum IV sonuclari yorumlar "
        "ve sinirlamalari tartisir. Bolum V sonuc ve icgoruleri ozetler.")

    # ---- II. YONTEM ----
    add_heading(doc, "II. YONTEM", size=10, small_caps=True)

    add_heading(doc, "A. Veri Seti ve Problem Tanimi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Kullanilan veri seti, 2014-2018 yillari icin ABD borsalarinda islem goren sirketlere "
        "ait yil basina yaklasik 4000 gozlem ve 224 finansal gosterge icermektedir [1]. "
        "Problemi regresyon cercevesine oturtabilmek icin ardisik yil ciftleri "
        "(2014-2015, 2015-2016, 2016-2017, 2017-2018) Ticker (hisse sembolu) anahtariyla "
        "ic birlesim (inner join) uygulanarak birlestirilmis; her ornek icin yil N gostergeleri "
        "ozellik, yil N+1 geliri ise hedef olarak atanmistir. Pozitif olmayan gelir degerleri "
        "filtrelenerek nihai veri setinde 15.497 gozlem elde edilmistir. Hedef degisken "
        "(Revenue_next) cok genis bir dinamik araligi (yuz binlerce dolardan yuz milyarlarca "
        "dolara kadar) icerdigi icin logaritmik donusum log1p uygulanmistir.")

    add_heading(doc, "B. Veri Temizleme",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Birlesik veri setinde yaklasik %5 oraninda eksik deger saptanmistir. "
        "Uygulanan temizleme adimlari:")

    items = [
        ("1) Yuksek eksiklikli sutunlarin kaldirilmasi:",
         " %40 esiginin uzerinde eksigi olan sutunlar cikarilmistir. Bu oranda eksiklik, "
         "atama ile guvenilirlik kaybi yaratmaktadir."),
        ("2) Sonsuz degerlerin islenmesi:",
         " Oran hesaplamalarinda sifira bolme nedeniyle olusan +/-inf degerleri NaN "
         "olarak yeniden isaretlenmistir."),
        ("3) Duplike satirlarin kaldirilmasi:",
         " Yinelenen satirlar veri setinden cikarilmistir."),
        ("4) Medyan atama:",
         " Kalan sayisal sutunlardaki eksik degerler medyan ile doldurulmustur. Medyan, "
         "finansal verilerdeki uc degerlere karsi ortalamadan daha dayaniklidir [2]."),
        ("5) Mod atama:",
         " Kategorik Sector sutunundaki olasi bos degerler en sik gozlemlenen deger (mod) "
         "ile doldurulmustur."),
    ]
    for head, body in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(head)
        r1.italic = True
        r1.font.size = Pt(10)
        r1.font.name = "Times New Roman"
        r2 = p.add_run(body)
        r2.font.size = Pt(10)
        r2.font.name = "Times New Roman"

    add_heading(doc, "C. Ozellik Muhendisligi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Sizinti kaynagi olabilecek Ticker, Class, yila ozgu PRICE VAR [%] sutunlari ve "
        "yil bilgisi iceren SourceYear ozellik matrisinden cikarilmistir. Kategorik Sector "
        "degiskeni One-Hot Encoding ile on bir ikili sutuna donusturulmus; boylece sektorler "
        "arasinda yapay bir siralama olusturulmamistir. Tum sayisal ozelliklere StandardScaler "
        "uygulanarak ortalama 0 ve standart sapma 1 olacak sekilde normalize edilmistir; bu "
        "adim dogrusal modeller ve KNN gibi olcek duyarli algoritmalar icin kritik oneme "
        "sahiptir [3]. Boyut lanetini azaltmak icin SelectKBest yontemi ve ANOVA F-istatistigi "
        "(f_regression) kullanilarak hedef ile en iliskili 50 ozellik secilmistir.")

    add_heading(doc, "D. Model Secimi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Farkli varsayimlarla calisan yedi regresyon modeli karsilastirilmistir: Dogrusal "
        "Regresyon baz cizgi olarak; Ridge (L2) ve Lasso (L1) duzenlilestirilmis dogrusal "
        "modeller olarak; Karar Agaci dogrusal olmayan ve yorumlanabilir; Rastgele Orman ve "
        "Gradient Boosting topluluk (ensemble) yontemleri olarak; K-En Yakin Komsu ise "
        "parametrik olmayan bir yaklasim olarak dahil edilmistir. Topluluk yontemleri "
        "finansal tahminlerde yaygin olarak ustun sonuclar vermektedir [4], [5]. Model "
        "hiperparametreleri literaturde yaygin degerlerle sabitlenmistir (orn. Rastgele "
        "Orman icin 300 agac ve derinlik 15; Gradient Boosting icin 300 agac, derinlik 4 "
        "ve ogrenme orani 0.05).")

    add_heading(doc, "E. Model Degerlendirmesi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Veri seti %80 egitim ve %20 test olmak uzere ayrilmistir. Egitim kumesi uzerinde "
        "5-katli capraz dogrulama uygulanmistir; bu yontem modelin farkli alt kumelerde "
        "tutarli performans sergilemesini ve asiri uyumdan (overfitting) kacinilmasini saglar. "
        "Degerlendirme metrikleri olarak R2, Ortalama Mutlak Hata (MAE), Kok Ortalama Kare "
        "Hata (RMSE) ve gercek olcekte MAPE hesaplanmistir. Metrikler hem log olceginde hem "
        "de expm1 ters donusumu sonrasinda gercek dolar olceginde raporlanmistir. Hedefin "
        "genis dinamik araligindan oturu log olcekli R2 ana karsilastirma metrigi olarak "
        "benimsenmistir [6].")

    add_figure(doc, os.path.join(OUT_DIR, "02_target_distribution.png"),
               "Sek. 1. Hedef degiskenin (Revenue_next) ham ve log1p donusumlu dagilimi.")

    # ---- III. SONUCLAR ----
    add_heading(doc, "III. SONUCLAR", size=10, small_caps=True)

    add_heading(doc, "A. Kesifsel Veri Analizi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Sek. 1, hedef degiskenin dagilimini gostermektedir. Ham olcekte veri saga dogru "
        "son derece carpik bir dagilim sergilerken, log1p donusumu sonrasinda cansimsal "
        "(normal) dagilima daha yakin bir gorunum almaktadir. Bu gozlem, logaritmik donusumun "
        "regresyon basarimi icin kritik olacagini on gostermektedir.")
    add_body(doc,
        "Sek. 2, sektorlere gore bir sonraki yil gelirinin kutu grafigini gostermektedir. "
        "Energy ve Consumer Defensive sektorlerinin medyan gelirleri ortalamanin uzerinde "
        "iken Financial Services, Technology ve Healthcare sektorleri genis bir dagilim "
        "sergilemektedir.")

    add_figure(doc, os.path.join(OUT_DIR, "03_sector_boxplot.png"),
               "Sek. 2. Sektorlere gore log1p(Revenue_next) kutu grafigi.")

    add_body(doc,
        "Sek. 3, kaynak yila gore ortalama hedef gelirini gostermektedir. 2014-2018 donemi "
        "boyunca ortalama gelirlerin benzer seviyede kaldigi gozlemlenmektedir; bu durum "
        "zaman baglantili bir biasin sinirli oldugunu dusundurur.")

    add_figure(doc, os.path.join(OUT_DIR, "04_yearly_mean.png"),
               "Sek. 3. Kaynak yila gore hedef gelirin ortalamasi.")

    add_body(doc,
        "Sek. 4, yil N geliri ile yil N+1 geliri arasindaki log-log olcekli iliskiyi "
        "gostermektedir. Veri noktalarinin y=x dogrusuna cok yakin dagilmasi, bir sirketin "
        "bu yilki gelirinin bir sonraki yil gelirinin en guclu habercisi oldugunu "
        "kanitlamaktadir. Log-log korelasyon katsayisi yaklasik 0.98 olarak hesaplanmistir.")

    add_figure(doc, os.path.join(OUT_DIR, "06_rev_scatter.png"),
               "Sek. 4. Yil N geliri ile yil N+1 geliri arasindaki log-log iliski.")

    add_body(doc,
        "Sek. 5, hedef ile en yuksek mutlak korelasyona sahip 15 ozellik arasindaki "
        "iliskileri gostermektedir. Beklenildigi gibi Revenue, Cost of Revenue, Gross Profit "
        "ve Operating Income gibi olcek ve karlilik metrikleri hedefle yuksek korelasyonludur.")

    add_figure(doc, os.path.join(OUT_DIR, "07_correlation_heatmap.png"),
               "Sek. 5. En iliskili 15 ozellik ve hedef icin korelasyon isi haritasi.")

    add_heading(doc, "B. Ozellik Onemi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Sek. 6, SelectKBest sonrasinda F-skoru en yuksek 20 ozelligi gostermektedir. "
        "Listenin ust siralarinda Revenue, Cost of Revenue, Gross Profit, Total Assets ve "
        "Operating Cash Flow gibi sirket buyuklugunu ve karliligini dogrudan temsil eden "
        "metrikler yer almaktadir. Bu gozlem, gelir tahmininde sirket olceginin birinci "
        "derecede belirleyici oldugunu gostermektedir.")

    add_figure(doc, os.path.join(OUT_DIR, "08_feature_importance.png"),
               "Sek. 6. SelectKBest ile en onemli 20 ozellik.")

    add_heading(doc, "C. Model Karsilastirmasi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Tablo I, yedi modelin test kumesi uzerindeki performansini log olcekte R2, MAE ve "
        "RMSE metrikleri ile ozetlemektedir. Gradient Boosting modeli 0.966 R2, 0.228 MAE ve "
        "0.467 RMSE degerleri ile en iyi sonucu vermistir. Rastgele Orman modeli cok yakin "
        "bir performansla ikinci siradadir (R2=0.964). Lineer modeller (Dogrusal, Ridge, Lasso) "
        "log-R2 bazinda 0.32-0.40 araliginda kalarak dogrusal olmayan yapiyi yakalayamadigini "
        "gostermistir.")

    # ---- Tablo I ----
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(0)
    cr1 = cap.add_run("TABLO I\n")
    cr1.bold = True
    cr1.font.size = Pt(9)
    cr1.font.name = "Times New Roman"
    cr2 = cap.add_run("MODEL KARSILASTIRMA SONUCLARI (LOG OLCEK)")
    cr2.italic = True
    cr2.font.size = Pt(9)
    cr2.font.name = "Times New Roman"

    headers = ["Model", "CV R2", "CV MAE", "CV RMSE", "R2", "MAE", "RMSE"]
    rows = [
        ["Gradient Boosting", "0.962", "0.228", "0.493", "0.966", "0.228", "0.467"],
        ["Rastgele Orman",    "0.962", "0.229", "0.491", "0.964", "0.233", "0.485"],
        ["Karar Agaci",       "0.945", "0.271", "0.590", "0.950", "0.262", "0.567"],
        ["KNN (k=7)",         "0.729", "0.918", "1.313", "0.733", "0.905", "1.312"],
        ["Dogrusal Regresyon","0.318", "1.510", "2.074", "0.399", "1.508", "1.968"],
        ["Ridge (alpha=1)",   "0.331", "1.508", "2.057", "0.399", "1.508", "1.969"],
        ["Lasso (alpha=0.01)","0.376", "1.510", "1.988", "0.380", "1.526", "2.000"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = "Times New Roman"
    # data rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Times New Roman"

    doc.add_paragraph()  # tablo sonrasi bosluk

    add_figure(doc, os.path.join(OUT_DIR, "09_model_comparison.png"),
               "Sek. 7. Modellerin test kumesindeki R2, MAE ve RMSE karsilastirmasi (log olcek).")

    add_body(doc,
        "Sek. 8, her model icin tahmin-gercek scatter plot kareleri gostermektedir. "
        "Gradient Boosting ve Rastgele Orman tahminleri y=x dogrusu etrafinda dar bir band "
        "icinde toplanirken, dogrusal modellerde uc deger tahminlerinde belirgin sapmalar "
        "gozlemlenmektedir.")

    add_figure(doc, os.path.join(OUT_DIR, "10_pred_vs_true.png"),
               "Sek. 8. Tum modeller icin tahmin vs gercek (log olcek).")

    add_heading(doc, "D. En Iyi Modelin Detayli Incelemesi",
                size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, bold=False,
                space_before=4, space_after=2)
    add_body(doc,
        "Sek. 9, en iyi model olan Gradient Boosting icin artik (residual) grafikleri "
        "gostermektedir. Artiklarin sifir etrafinda simetrik olarak dagilmasi ve dar bir "
        "araliga sahip olmasi modelin yanli olmadigini ve hatalarin gurultu gibi "
        "davrandigini dogrulamaktadir.")

    add_figure(doc, os.path.join(OUT_DIR, "11_residuals.png"),
               "Sek. 9. Gradient Boosting modeli icin artik dagilimi ve histogrami.")

    add_body(doc,
        "Sek. 10, en iyi modelin ogrendigi en onemli 20 ozelligi gostermektedir. Revenue "
        "tek basina ozellik oneminin buyuk bolumunu aciklamaktadir; bunu Cost of Revenue, "
        "Operating Expenses ve Total Assets izlemektedir.")

    add_figure(doc, os.path.join(OUT_DIR, "12_best_model_importance.png"),
               "Sek. 10. Gradient Boosting modelinin en onemli 20 ozelligi.")

    # ---- IV. TARTISMA ----
    add_heading(doc, "IV. TARTISMA", size=10, small_caps=True)
    add_body(doc,
        "Elde edilen sonuclar, finansal tablo verilerinin kisa vadeli gelir tahmini icin "
        "cok guclu bir sinyal tasidigini gostermektedir. Gradient Boosting ve Rastgele Orman "
        "modelleri log olcekte 0.96 uzerinde R2 degerine ulasmis; bu sonuc, topluluk tabanli "
        "agac yontemlerinin finansal verilerdeki dogrusal olmayan iliskiler ve uc degerler "
        "karsisinda ustun oldugunu dogrulamaktadir.")
    add_body(doc,
        "Dogrusal modellerin dusuk R2 skorlari iki nedenle aciklanabilir. Birincisi, "
        "gelir-maliyet iliskisi katlamalidir ve log donusumune ragmen tam "
        "dogrusallasmayabilir. Ikincisi, coklu dogrusallik (multicollinearity) yuksek "
        "oldugundan, L2 ve L1 duzenlilestirme tek basina yeterli olmamistir. KNN modeli ise "
        "yerel komsuluk etkisiyle orta seviye R2 degerine ulasirken, yuksek boyutta "
        "etkinligini kaybetmistir.")
    add_body(doc,
        "Ozellik onem analizleri, yil N gelirinin yil N+1 gelirini aciklamakta neredeyse "
        "tek basina yeterli oldugunu ortaya koymaktadir. Bu bulgu gelir otokorelasyonunun "
        "finansal tahminde ne kadar guclu oldugunu dogrulamaktadir. Bununla birlikte, "
        "Rastgele Orman ve Gradient Boosting gibi modeller yalnizca Revenue sutununa bagli "
        "kalmamis, karlilik, maliyet ve olcek metriklerini de birlestirerek ek dogruluk "
        "kazanimi saglamistir.")
    add_body(doc,
        "Calismanin baslica sinirlamalari sunlardir: (i) Yalnizca finansal tablo "
        "gostergeleri kullanilmis, makroekonomik degiskenler, haber/duygu analizi veya "
        "sektor trendleri dahil edilmemistir. (ii) Model sadece 2014-2018 donemine ait "
        "verilerle egitilmis olup daha yeni donemlerde guvenilirligi dogrulanmaya "
        "muhtactir. (iii) Model ardisik iki yil arasinda bir eslestirme yaparken uzun "
        "vadeli trendleri dogrudan kullanmamaktadir.")
    add_body(doc,
        "Eyleme donuk icgoruler acisindan, gelir tahmini sirketlerin iceride butce "
        "planlamasinda ve disarida yatirimci degerlendirmesinde hizli bir referans araci "
        "olarak kullanilabilir. Gradient Boosting gibi kutu-ici modellerin yuksek dogrulugu, "
        "kural tabanli geleneksel tahmin yontemlerini tamamlayici bir otomatik tahmin "
        "katmani sunmaktadir.")

    # ---- V. SONUC ----
    add_heading(doc, "V. SONUC", size=10, small_caps=True)
    add_body(doc,
        "Bu calismada, 2014-2018 yillarina ait ABD borsasi finansal verileri kullanilarak "
        "sirketlerin bir sonraki yilki gelirlerinin regresyon yontemleriyle tahmin edildigi "
        "uctan uca bir veri madenciligi sureci uygulanmistir. Ardisik yillarin "
        "birlestirilmesi, kapsamli veri temizleme, log donusumu, one-hot kodlama, "
        "standartlastirma ve ANOVA tabanli ozellik secimi adimlarinin ardindan yedi farkli "
        "model 5-katli capraz dogrulama ile karsilastirilmistir. Gradient Boosting modeli "
        "log olcekli test kumesinde 0.966 R2, 0.228 MAE ve 0.467 RMSE degerleriyle en iyi "
        "performansi saglamistir.")
    add_body(doc,
        "Gelecek calismalarda zaman serisi yaklasimlari (ARIMA, LSTM, Temporal Fusion "
        "Transformer), makroekonomik degiskenler, sektor trendleri ve duygu/haber analizinin "
        "dahil edilmesi modelin genelleme yetenegini artirabilir. Ayrica hiperparametre "
        "optimizasyonu (Grid Search, Bayesian Optimization) ve SHAP gibi aciklanabilirlik "
        "yontemleri modelin karar surecini daha seffaf hale getirebilir.")

    # ---- KAYNAKLAR ----
    add_heading(doc, "KAYNAKLAR", size=10, small_caps=True,
                align=WD_ALIGN_PARAGRAPH.CENTER)
    refs = [
        "[1] N. Carbone, \"200+ Financial Indicators of US stocks (2014-2018),\" Kaggle, 2019. "
        "[Cevrimici]. Mevcut: https://www.kaggle.com/datasets/cnic92/200-financial-indicators-of-us-stocks-20142018",
        "[2] J. Han, M. Kamber, and J. Pei, Data Mining: Concepts and Techniques, 3rd ed. "
        "Waltham, MA, USA: Morgan Kaufmann, 2011.",
        "[3] F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" "
        "J. Mach. Learn. Res., vol. 12, pp. 2825-2830, Oct. 2011.",
        "[4] L. Breiman, \"Random forests,\" Mach. Learn., vol. 45, no. 1, pp. 5-32, 2001.",
        "[5] J. H. Friedman, \"Greedy function approximation: A gradient boosting machine,\" "
        "Ann. Statist., vol. 29, no. 5, pp. 1189-1232, 2001.",
        "[6] T. Hastie, R. Tibshirani, and J. Friedman, The Elements of Statistical Learning, "
        "2nd ed. New York, NY, USA: Springer, 2009.",
        "[7] G. James, D. Witten, T. Hastie, and R. Tibshirani, An Introduction to Statistical "
        "Learning, 2nd ed. New York, NY, USA: Springer, 2021.",
        "[8] E. F. Fama and K. R. French, \"Common risk factors in the returns on stocks and "
        "bonds,\" J. Financ. Econ., vol. 33, no. 1, pp. 3-56, 1993.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ref)
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"

    doc.save(DOCX_PATH)
    print(f"OK: {DOCX_PATH}")


if __name__ == "__main__":
    build()
